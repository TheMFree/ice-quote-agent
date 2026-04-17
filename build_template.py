"""
Generates templates/ICE_Contractors_Proposal_Template.docx from scratch.

This avoids shipping a binary .docx in the repo and lets us rebuild the
template on the VPS during bootstrap. Run:

    python build_template.py

Output: templates/ICE_Contractors_Proposal_Template.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Twips

# ---- constants --------------------------------------------------------------
FONT = "Arial"
GRAY = "CCCCCC"
LIGHT_BLUE = "DCE6F1"
PAGE_WIDTH_TWIPS = 12240        # 8.5"
PAGE_HEIGHT_TWIPS = 15840       # 11"


# ---- low level helpers ------------------------------------------------------
def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_cell_border(cell, style="single", size=4, color=GRAY,
                     sides=("top", "bottom", "left", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side in sides:
        el = tcBorders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tcBorders.append(el)
        el.set(qn("w:val"), style)
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def _no_borders(cell):
    _set_cell_border(cell, style="nil", size=0, color="FFFFFF")


def _set_cell_width(cell, twips):
    tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW()
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(twips))


def _set_cell_valign(cell, vert):
    tcPr = cell._tc.get_or_add_tcPr()
    el = tcPr.find(qn("w:vAlign"))
    if el is None:
        el = OxmlElement("w:vAlign")
        tcPr.append(el)
    el.set(qn("w:val"), vert)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _apply_run(run, *, bold=False, italic=False, size=11, font=FONT, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:eastAsia"), font)


def _add_paragraph(container, text="", *, bold=False, italic=False, size=11,
                   align=None, space_before=None, space_after=80):
    p = container.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _apply_run(run, bold=bold, italic=italic, size=size)
    return p


def _add_cell_paragraph(cell, text="", *, bold=False, italic=False, size=11,
                        align=None, space_after=4):
    if (cell.paragraphs and not cell.paragraphs[0].text
            and not getattr(cell, "_first_used", False)):
        p = cell.paragraphs[0]
        cell._first_used = True
    else:
        p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        _apply_run(run, bold=bold, italic=italic, size=size)
    return p


def _add_bullet(container, text="", *, size=11):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Twips(720)
    pf.first_line_indent = Twips(-288)
    pf.space_after = Pt(3)
    run = p.add_run("\u2022  " + (text or ""))
    _apply_run(run, size=size)
    return p


BLANK_UNDERLINE = "______________________________________________________________________"


def _add_blank_bullet(container):
    return _add_bullet(container, BLANK_UNDERLINE)


# ---- footer with page numbers ----------------------------------------------
def _set_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = p.add_run("ICE Contractors, Inc \u2014 Proposal    |    Page ")
    _apply_run(run1, italic=True, size=9)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    r_page = p.add_run()
    _apply_run(r_page, italic=True, size=9)
    r_page._r.append(fldChar1)
    r_page._r.append(instrText)
    r_page._r.append(fldChar2)

    r_sep = p.add_run(" of ")
    _apply_run(r_sep, italic=True, size=9)

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText")
    instrText2.set(qn("xml:space"), "preserve")
    instrText2.text = " NUMPAGES "
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    r_total = p.add_run()
    _apply_run(r_total, italic=True, size=9)
    r_total._r.append(fldChar3)
    r_total._r.append(instrText2)
    r_total._r.append(fldChar4)


# ---- standard content blocks ------------------------------------------------
TERMS_BULLETS = [
    "Payment terms to be progress payments with final payment upon completion and acceptance of the Work.",
    "Price based on day shift only. Any required overtime or shift work beyond this schedule will be addressed via change order.",
    "Quote based on having access to install the I&E portion of the work after mechanical equipment is set and needed work areas are available.",
    "Any required rock blasting for excavation shall be provided by others.",
    "Major equipment provided by owner shall be set in place by others.",
    "Construction mats to be provided by others, if applicable.",
    "10% will be billed ARO (at receipt of order).",
    "Mobilization billed at receipt of order and upon notice to proceed.",
    "Price does not include temporary power. Adequate site power to be provided by owner.",
    "Quote assumes one mobilization and demobilization. Any deviation not the fault of ICE Contractors will be offered at an additional charge.",
    "Specialty and long-lead materials to be billed upon delivery at the next applicable progress billing.",
    "Delays or efficiency loss incurred due to circumstances outside our control will be billed at our current T&M rates.",
    "Any production escalations required through no fault of ICE Contractors will be offered at additional cost on a situational basis.",
    "Structural support and welding provided by mechanical contractor.",
    "Red-lined as-built drawings to be provided for owner markup. Final CAD drawings provided by others.",
    "Contractor shall comply with DOT Anti-Drug, Alcohol Misuse, and OQ Training requirements per bid instructions.",
    "All work per NEC 2023, NFPA 70, and applicable client/owner specifications.",
    "Rain-out or site shutdown days shall be added to the original mechanical completion date on a day-for-day basis.",
]

EXCLUSIONS_BULLETS = [
    "Owner-furnished equipment procurement.",
    "DCS/SCADA programming and configuration.",
    "Fire and gas detection system (if separate from I&E scope).",
    "Painting beyond NEC requirements.",
    "Permit fees and engineering stamps.",
    "Structural modifications or civil/concrete work (foundations and piers provided by others).",
    "Cathodic protection system.",
    "Communication tower or antenna work.",
    "Environmental remediation or hazmat abatement.",
    "Fencing modifications beyond grounding connections.",
]

GENERAL_NOTES_BULLETS = [
    "Terms to be 2%10, Net 30.",
    "Price based on day shift only.",
    "Price is based on a _____ schedule (e.g., 4-10s, 5-8s, 6-10s).",
    "Quote based on having time to install the I&E portion of the work after mechanical equipment and needed work areas are available.",
    "Any required rock blasting for excavation shall be provided by others.",
    "Major equipment provided by customer shall be set in place by others.",
    "If applicable, construction mats to be provided by others.",
    "Payment terms to be Net 30 with progress billings performed bi-weekly based on agreement of work complete.",
    "10% will be billed ARO.",
    "Price does not include temporary power.",
    "Any required concrete testing provided by others.",
    "Specialty materials to be billed upon delivery at the next applicable progress billing.",
    "Delays/efficiency loss incurred due to circumstances out of our control will be billed at our current T&M rates.",
    "Any production escalations required at no fault of ICE Contractors will be offered at an additional cost and quoted on a situational basis.",
    "Mobilization billed at receipt of order and upon notice to proceed.",
    "Quote assumes one mobilization and demobilization. Any deviation that is not the fault of ICE Contractors will be offered at an additional charge.",
    "Structural support and welding provided by mechanical contractor.",
    "Red-lined drawings to be provided for owner mark-up. Final CAD drawings provided by others.",
]


# ---- high level section builders --------------------------------------------
def build_header_table(doc):
    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    widths = [2340, 4680, 2340]
    cells = t.rows[0].cells
    for cell, w in zip(cells, widths):
        _set_cell_width(cell, w)
        _no_borders(cell)
        _set_cell_margins(cell)

    _add_cell_paragraph(cells[0], "ICE", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_cell_paragraph(cells[0], "CONTRACTORS", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    _add_cell_paragraph(cells[1], "ICE CONTRACTORS, INC", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_cell_paragraph(cells[1], "Instrumentation & Electrical Contractors", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_cell_paragraph(cells[1], "FEIN: 35-2206612", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    _add_cell_paragraph(cells[2], "INSTRUMENTATION", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_cell_paragraph(cells[2], "& ELECTRICAL", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_cell_paragraph(cells[2], "CONTRACTORS, INC.", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    return t


def _add_project_info_row(table, label, value):
    row = table.add_row()
    c1, c2 = row.cells
    _set_cell_width(c1, 2340)
    _set_cell_width(c2, 7020)
    _shade(c1, LIGHT_BLUE)
    _set_cell_border(c1)
    _set_cell_border(c2)
    _set_cell_margins(c1)
    _set_cell_margins(c2)
    _add_cell_paragraph(c1, label, bold=True, size=11)
    _add_cell_paragraph(c2, value, size=11)


def build_project_info_table(doc):
    t = doc.add_table(rows=0, cols=2)
    t.autofit = False
    rows = [
        ("Project:", ""),
        ("Agreement No.:", ""),
        ("WBS #:", ""),
        ("Client:", ""),
        ("Location:", ""),
        ("Proposal Date:", ""),
        ("Quote No.:", ""),
        ("Prepared By:", "Michael G. Freeman, Jr. \u2014 President"),
        ("Valid For:", "30 Days from Proposal Date"),
    ]
    for label, value in rows:
        _add_project_info_row(t, label, value)
    return t


def build_pricing_table(doc):
    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    hdr = t.rows[0].cells
    widths = [780, 5580, 3000]
    for cell, w in zip(hdr, widths):
        _set_cell_width(cell, w)
        _shade(cell, LIGHT_BLUE)
        _set_cell_border(cell)
        _set_cell_margins(cell)
    _add_cell_paragraph(hdr[0], "#", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_cell_paragraph(hdr[1], "Description", bold=True, size=11)
    _add_cell_paragraph(hdr[2], "Amount", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)

    def _row(n, desc, amount, bold=False):
        row = t.add_row()
        cells = row.cells
        for cell, w in zip(cells, widths):
            _set_cell_width(cell, w)
            _set_cell_border(cell)
            _set_cell_margins(cell)
        _add_cell_paragraph(cells[0], n, bold=bold, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        _add_cell_paragraph(cells[1], desc, bold=bold, size=11)
        _add_cell_paragraph(cells[2], amount, bold=bold, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _row("1", "Material", "$________________")
    _row("2", "Labor & Equipment", "$________________")
    _row("", "TOTAL LUMP SUM PRICE", "$________________", bold=True)
    return t


def build_crew_table(doc):
    t = doc.add_table(rows=0, cols=2)
    t.autofit = False
    widths = [2880, 6480]
    rows = [
        ("Total Man-Hours:", ""),
        ("Crew Size:", ""),
        ("Estimated Duration:", ""),
        ("Work Schedule:", ""),
    ]
    for label, value in rows:
        row = t.add_row()
        c1, c2 = row.cells
        _set_cell_width(c1, widths[0])
        _set_cell_width(c2, widths[1])
        _shade(c1, LIGHT_BLUE)
        for cell in (c1, c2):
            _set_cell_border(cell)
            _set_cell_margins(cell)
        _add_cell_paragraph(c1, label, bold=True, size=11)
        _add_cell_paragraph(c2, value, size=11)
    return t


def build_auth_table(doc):
    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    widths = [4380, 600, 4380]
    cells = t.rows[0].cells
    for cell, w in zip(cells, widths):
        _set_cell_width(cell, w)
        _set_cell_margins(cell)
        _no_borders(cell)
        _set_cell_valign(cell, "top")

    left_lines = [
        "_____________________________________",
        "Accepted by (Client Signature)",
        "",
        "_____________________________________",
        "Print Name / Title",
        "",
        "_____________________________________",
        "Date",
    ]
    right_lines = [
        "_____________________________________",
        "Michael G. Freeman, Jr. \u2014 President",
        "ICE Contractors, Inc",
        "",
        "_____________________________________",
        "Print Name / Title",
        "",
        "_____________________________________",
        "Date",
    ]
    for line in left_lines:
        _add_cell_paragraph(cells[0], line, size=11, space_after=14)
    for line in right_lines:
        _add_cell_paragraph(cells[2], line, size=11, space_after=14)
    return t


def h1(doc, text):
    _add_paragraph(doc, text, bold=True, size=12, space_before=11, space_after=6)


def h2(doc, text):
    _add_paragraph(doc, text, bold=True, size=11, space_before=8, space_after=4)


# ---- main assembly ----------------------------------------------------------
def build(path: Path):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Twips(PAGE_WIDTH_TWIPS)
    section.page_height = Twips(PAGE_HEIGHT_TWIPS)
    section.top_margin = Twips(1080)
    section.bottom_margin = Twips(1080)
    section.left_margin = Twips(1080)
    section.right_margin = Twips(1080)
    section.orientation = WD_ORIENT.PORTRAIT

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)

    _set_footer(doc)

    build_header_table(doc)
    _add_paragraph(doc, "")

    _add_paragraph(doc, "LUMP SUM PROPOSAL", bold=True, size=15,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)

    h1(doc, "PROJECT INFORMATION")
    build_project_info_table(doc)
    _add_paragraph(doc, "")

    h1(doc, "SCOPE OF WORK")
    _add_paragraph(doc,
        "ICE Contractors, Inc. proposes to furnish all labor, materials, equipment, "
        "tools, and supervision required to complete the instrumentation and electrical "
        "(I&E) installation for the ________________________________ project at "
        "________________________________. Work is per [Drawing Reference / Specification "
        "Reference], the Cable Schedule, P&IDs, Plot Plan, and all applicable "
        "client/owner specifications.",
        size=11)
    _add_paragraph(doc, "Work includes, but is not limited to:", size=11,
                   space_before=4, space_after=4)
    for _ in range(8):
        _add_blank_bullet(doc)
    _add_paragraph(doc, "")

    h1(doc, "LUMP SUM PRICE")
    build_pricing_table(doc)
    _add_paragraph(doc, "")

    h1(doc, "ESTIMATED CREW & SCHEDULE")
    build_crew_table(doc)
    _add_paragraph(doc, "")

    h1(doc, "TERMS & CONDITIONS")
    _add_paragraph(doc,
        "This proposal is valid for 30 days. Prices are based on current material costs "
        "and may be subject to adjustment for market fluctuations. Changes to scope of "
        "work will be addressed via written change order.",
        italic=True, size=11, space_after=6)
    for b in TERMS_BULLETS:
        _add_bullet(doc, b, size=11)
    _add_paragraph(doc, "")

    h1(doc, "EXCLUSIONS")
    for b in EXCLUSIONS_BULLETS:
        _add_bullet(doc, b, size=11)

    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)

    h1(doc, "ATTACHMENT A \u2014 SCOPE OF WORK INCLUDED")
    _add_paragraph(doc,
        "The following is a summary of electrical and instrumentation scope included in "
        "this proposal. All work per NEC 2023 and applicable local codes. Material "
        "quantities derived from drawing-verified takeoff.",
        size=11)
    _add_paragraph(doc, "")

    subsections = [
        ("Service Entrance & Main Distribution", 3),
        ("Sub-Distribution", 3),
        ("Conduit & Cable Tray", 3),
        ("Wire & Cable", 3),
        ("Equipment Connections", 3),
        ("Instrumentation", 3),
        ("Lighting", 2),
        ("Grounding", 2),
        ("Testing & Commissioning", 2),
        ("Long-Lead Items", 3),
        ("Assumptions", 3),
        ("Pending Clarifications", 2),
    ]
    for title, blanks in subsections:
        h2(doc, title)
        for _ in range(blanks):
            _add_blank_bullet(doc)
        _add_paragraph(doc, "")

    page2 = doc.add_paragraph()
    page2.add_run().add_break(WD_BREAK.PAGE)

    h1(doc, "ATTACHMENT B \u2014 GENERAL NOTES")
    for b in GENERAL_NOTES_BULLETS:
        _add_bullet(doc, b, size=11)
    _add_paragraph(doc, "")

    h1(doc, "AUTHORIZATION")
    _add_paragraph(doc,
        "Acceptance of this proposal constitutes agreement to the terms, conditions, "
        "exclusions, and general notes contained herein.",
        italic=True, size=11)
    _add_paragraph(doc, "")
    build_auth_table(doc)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    size = path.stat().st_size
    print(f"Wrote {path} ({size} bytes)")


if __name__ == "__main__":
    out = Path(__file__).resolve().parent / "templates" / "ICE_Contractors_Proposal_Template.docx"
    build(out)
