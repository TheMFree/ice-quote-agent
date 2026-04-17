"""Word document parser."""
from __future__ import annotations
import logging
import shutil
import subprocess
import tempfile
from pathlib import Path

log = logging.getLogger("ice_quote_agent")


def _convert_doc_to_docx(src: Path) -> Path | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        log.warning("soffice/libreoffice not found; cannot convert .doc")
        return None
    tmp = Path(tempfile.mkdtemp())
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "docx",
             "--outdir", str(tmp), str(src)],
            check=True, capture_output=True, timeout=60,
        )
        out = tmp / (src.stem + ".docx")
        return out if out.exists() else None
    except Exception as e:
        log.exception("soffice conversion failed for %s: %s", src.name, e)
        return None


def parse(path: Path, out) -> None:
    target = path
    if path.suffix.lower() == ".doc":
        converted = _convert_doc_to_docx(path)
        if converted is None:
            out.add_text(f"[Could not convert legacy .doc file {path.name}]",
                         source_label=path.name)
            return
        target = converted

    try:
        import docx
    except ImportError:
        out.add_text(f"[python-docx not installed; skipping {path.name}]",
                     source_label=path.name)
        return

    d = docx.Document(str(target))
    parts: list[str] = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    if parts:
        out.add_text(f"[Word doc {path.name}]\n" + "\n".join(parts),
                     source_label=path.name)
