"""Fills the ICE proposal template with QuoteData."""
from __future__ import annotations
import copy
import logging
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document
from docx.document import Document as _Doc
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .schema import QuoteData, ScopeCategory

log = logging.getLogger("ice_quote_agent")


def _iter_block_items(parent):
    body = parent.element.body if isinstance(parent, _Doc) else parent._element
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _set_cell_text(cell: _Cell, text: str, bold: bool = False) -> None:
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(11)


def _find_label_cell(table: Table, label: str) -> Optional[int]:
    needle = label.strip().lower().rstrip(":")
    for i, row in enumerate(table.rows):
        first = row.cells[0].text.strip().lower().rstrip(":")
        if first == needle or first.startswith(needle):
            return i
    return None


def _fill_kv_table(table: Table, mapping: dict[str, Optional[str]]) -> None:
    for label, value in mapping.items():
        if value is None or value == "":
            continue
        idx = _find_label_cell(table, label)
        if idx is None:
            continue
        _set_cell_text(table.rows[idx].cells[1], value)


def fill_template(template_path: Path, data: QuoteData, output_path: Path) -> Path:
    doc = Document(str(template_path))

    if data.document_type:
        _replace_title(doc, data.document_type)

    for table in doc.tables:
        first_col_labels = {row.cells[0].text.strip().lower() for row in table.rows}

        if "project:" in first_col_labels and "client:" in first_col_labels:
            _fill_kv_table(table, {
                "Project": data.project,
                "Agreement No.": data.agreement_no,
                "WBS #": data.wbs_no,
                "Client": data.client,
                "Owner\u2019s Rep": data.owner_rep,
                "Location": data.location,
                "Proposal Date": data.proposal_date,
                "Quote No.": data.quote_no,
                "Prepared By": data.prepared_by,
                "Valid For": data.valid_for,
            })
        elif "total man-hours:" in first_col_labels:
            _fill_kv_table(table, {
                "Total Man-Hours": data.total_man_hours,
                "Crew Size": data.crew_size,
                "Estimated Duration": data.estimated_duration,
                "Work Schedule": data.work_schedule,
            })
        elif _is_pricing_table(table):
            _fill_pricing_table(table, data)

    _fill_scope_section(doc, data)
    _fill_scope_categories(doc, data)
    _fill_named_list(doc, "Long-Lead Items", data.long_lead_items)
    _fill_named_list(doc, "Assumptions", data.assumptions)
    _fill_named_list(doc, "Pending Clarifications", data.pending_clarifications)

    # EXCLUSIONS is empty-by-default in the template.
    # If Claude returns any, create bullets from scratch under the heading.
    if data.additional_exclusions:
        _insert_bullets_under_heading(doc, "EXCLUSIONS", data.additional_exclusions)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.info("Saved filled template to %s", output_path)
    return output_path


def _replace_title(doc: _Doc, title: str) -> None:
    for p in doc.paragraphs:
        if p.text.strip().upper() in ("LUMP SUM PROPOSAL", "QUOTATION"):
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = title.upper()
            else:
                p.add_run(title.upper()).bold = True
            break


def _is_pricing_table(table: Table) -> bool:
    if not table.rows:
        return False
    headers = [c.text.strip().lower() for c in table.rows[0].cells]
    return "description" in headers and "amount" in headers


def _fill_pricing_table(table: Table, data: QuoteData) -> None:
    """Fill the Amount column in the pricing table.

    The pricing table layout is: [#, Description, Amount].
    Row 0 is the header, rows 1/2 are Material / Labor & Equipment,
    row 3 is TOTAL.
    """
    headers = [c.text.strip().lower() for c in table.rows[0].cells]
    try:
        amt = headers.index("amount")
    except ValueError:
        return

    if len(table.rows) > 1 and data.material_amount:
        _set_cell_text(table.rows[1].cells[amt], data.material_amount)
    if len(table.rows) > 2 and data.labor_equipment_amount:
        _set_cell_text(table.rows[2].cells[amt], data.labor_equipment_amount)
    if len(table.rows) > 3 and data.total_amount:
        _set_cell_text(table.rows[3].cells[amt], data.total_amount, bold=True)


def _fill_scope_section(doc: _Doc, data: QuoteData) -> None:
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        if p.text.strip().startswith("ICE Contractors, Inc") and data.scope_intro:
            _replace_paragraph_text(p, data.scope_intro)
            break

    blank_bullets: List[Paragraph] = []
    in_scope = False
    for p in paragraphs:
        txt = p.text.strip()
        up = txt.upper()
        if up == "SCOPE OF WORK":
            in_scope = True
            blank_bullets = []
            continue
        if in_scope and up in ("LUMP SUM PRICE", "PRICING", "ESTIMATED CREW & SCHEDULE"):
            _populate_bullets(blank_bullets, data.scope_bullets)
            in_scope = False
            blank_bullets = []
            break
        if in_scope and txt.startswith("___"):
            blank_bullets.append(p)


def _replace_paragraph_text(p: Paragraph, new_text: str) -> None:
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


def _populate_bullets(blank_paragraphs: List[Paragraph], values: List[str]) -> None:
    """Populate '___' placeholder bullets with values.

    Behavior:
      - If values == blanks: one-to-one replacement.
      - If values > blanks: replace all blanks, clone the last blank to fit the rest.
      - If values < blanks: replace the first N, DELETE leftover placeholders
        (prevents stray "___" lines in the final draft).
    """
    n = min(len(blank_paragraphs), len(values))
    for idx in range(n):
        _replace_paragraph_text(blank_paragraphs[idx], values[idx])

    if len(values) > len(blank_paragraphs) and blank_paragraphs:
        last = blank_paragraphs[-1]
        for val in values[len(blank_paragraphs):]:
            new_p = copy.deepcopy(last._element)
            last._element.addnext(new_p)
            new_para = Paragraph(new_p, last._parent)
            _replace_paragraph_text(new_para, val)
            last = new_para
    elif len(values) < len(blank_paragraphs):
        # Delete unused "___" placeholders so they don't appear in the draft.
        for p in blank_paragraphs[len(values):]:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)


def _fill_scope_categories(doc: _Doc, data: QuoteData) -> None:
    if not data.scope_categories:
        return
    by_heading = {c.heading.strip().lower(): c for c in data.scope_categories}
    paragraphs = doc.paragraphs
    current_heading: Optional[str] = None
    current_blanks: List[Paragraph] = []

    def _flush():
        if current_heading and current_heading in by_heading and current_blanks:
            _populate_bullets(current_blanks, by_heading[current_heading].items)

    in_attachment_a = False
    for p in paragraphs:
        txt = p.text.strip()
        up = txt.upper()
        if up.startswith("ATTACHMENT A"):
            in_attachment_a = True
            continue
        if up.startswith("ATTACHMENT B") or up == "AUTHORIZATION":
            _flush()
            in_attachment_a = False
            break
        if not in_attachment_a:
            continue

        is_heading = (
            txt and not txt.startswith("___")
            and len(txt) < 70
            and (p.runs and p.runs[0].bold)
        )
        if is_heading and up not in ("ATTACHMENT A — SCOPE OF WORK INCLUDED",):
            _flush()
            current_heading = txt.lower()
            current_blanks = []
            continue

        if txt.startswith("___") and current_heading:
            current_blanks.append(p)

    _flush()


def _fill_named_list(doc: _Doc, heading: str, items: List[str]) -> None:
    if not items:
        return
    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        if paragraphs[i].text.strip().lower() == heading.strip().lower():
            blanks: List[Paragraph] = []
            j = i + 1
            while j < len(paragraphs):
                t = paragraphs[j].text.strip()
                if t.startswith("___"):
                    blanks.append(paragraphs[j])
                elif t == "":
                    pass
                else:
                    break
                j += 1
            _populate_bullets(blanks, items)
            return
        i += 1


def _is_list_bullet(p: Paragraph) -> bool:
    """Return True if this paragraph uses list numbering (bullets)."""
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


def _find_template_bullet(doc: _Doc) -> Optional[Paragraph]:
    """Find any bullet paragraph we can clone to inject new ones."""
    for p in doc.paragraphs:
        if _is_list_bullet(p) and p.text.strip():
            return p
    return None


def _insert_bullets_under_heading(doc: _Doc, heading: str, items: List[str]) -> None:
    """Insert bullet paragraphs right after `heading`, cloning an existing
    bullet's formatting. Safe to call when the section has zero existing bullets.
    """
    if not items:
        return
    paragraphs = doc.paragraphs
    heading_p: Optional[Paragraph] = None
    heading_idx = -1
    for i, p in enumerate(paragraphs):
        if p.text.strip().upper() == heading.upper():
            heading_p = p
            heading_idx = i
            break
    if heading_p is None:
        return

    # Prefer the last existing bullet inside this section (for consistent formatting),
    # otherwise fall back to any bullet elsewhere in the doc (T&C typically).
    section_bullet: Optional[Paragraph] = None
    for j in range(heading_idx + 1, len(paragraphs)):
        t = paragraphs[j].text.strip()
        up = t.upper()
        if up in ("AUTHORIZATION",) or up.startswith("ATTACHMENT "):
            break
        if _is_list_bullet(paragraphs[j]):
            section_bullet = paragraphs[j]

    template_bullet = section_bullet or _find_template_bullet(doc)
    if template_bullet is None:
        # Last resort: create plain paragraphs after the heading.
        insert_after = heading_p
        for val in items:
            new_para = insert_after._parent.add_paragraph(val)
            # python-docx's add_paragraph appends to end of parent — reposition it
            insert_after._element.addnext(new_para._element)
            insert_after = new_para
        return

    insert_after = heading_p
    for val in items:
        new_p = copy.deepcopy(template_bullet._element)
        insert_after._element.addnext(new_p)
        new_para = Paragraph(new_p, insert_after._parent)
        _replace_paragraph_text(new_para, val)
        insert_after = new_para


def _append_to_named_list(doc: _Doc, heading: str, items: List[str]) -> None:
    """Back-compat shim — delegates to the new _insert_bullets_under_heading
    so empty sections (like the new EXCLUSIONS) still get bullets created.
    """
    _insert_bullets_under_heading(doc, heading, items)
